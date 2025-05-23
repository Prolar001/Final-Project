#!/usr/bin/env python3
"""
CSU CCIS Blog Word Documentation Generator
Creates an editable Word document with comprehensive project documentation
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

def create_word_documentation():
    """Generate comprehensive Word documentation"""
    
    # Create output directory
    os.makedirs('output', exist_ok=True)
    
    # Create new Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Define CSU CCIS colors
    csu_brown = RGBColor(139, 69, 19)  # Brown
    csu_green = RGBColor(34, 139, 34)  # Forest Green
    csu_gold = RGBColor(255, 215, 0)   # Gold
    
    # Title Page
    title = doc.add_heading('CSU CCIS Blog Application', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = csu_brown
    title_run.font.size = Pt(24)
    
    subtitle = doc.add_heading('Complete Project Documentation', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.color.rgb = csu_green
    
    doc.add_paragraph()
    
    # Institution info
    institution = doc.add_paragraph('Colorado State University', style='Normal')
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    college = doc.add_paragraph('College of Computing and Information Sciences', style='Normal')
    college.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}', style='Normal')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Page break
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.runs[0].font.color.rgb = csu_brown
    
    toc_items = [
        "1. Project Overview",
        "2. Entity-Relationship Diagram", 
        "3. Database Schema",
        "4. API Specification",
        "5. Authentication Flow",
        "6. Setup Guide",
        "7. Usage Guide",
        "8. Code Structure",
        "9. Technologies Used"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # Add documentation sections
    doc_sections = [
        ("README.md", "Project Overview"),
        ("docs/ERD.md", "Entity-Relationship Diagram"),
        ("docs/DATABASE_SCHEMA.md", "Database Schema"),
        ("docs/API_SPEC.md", "API Specification"),
        ("docs/AUTHENTICATION_FLOW.md", "Authentication Flow"),
        ("docs/SETUP_GUIDE.md", "Setup Guide"),
        ("docs/USAGE_GUIDE.md", "Usage Guide")
    ]
    
    for file_path, section_title in doc_sections:
        # Section heading
        heading = doc.add_heading(section_title, level=1)
        heading.runs[0].font.color.rgb = csu_brown
        
        if os.path.exists(file_path):
            # Read file content
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Process markdown content
            lines = content.split('\n')
            in_code_block = False
            
            for line in lines:
                line = line.strip()
                
                if line.startswith('```'):
                    in_code_block = not in_code_block
                    continue
                
                if in_code_block:
                    # Add code as monospace
                    code_para = doc.add_paragraph(line)
                    if code_para.runs:
                        code_run = code_para.runs[0]
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(9)
                    continue
                
                if line.startswith('# '):
                    # Main heading
                    h = doc.add_heading(line[2:], level=2)
                    h.runs[0].font.color.rgb = csu_green
                elif line.startswith('## '):
                    # Sub heading
                    h = doc.add_heading(line[3:], level=3)
                    h.runs[0].font.color.rgb = csu_green
                elif line.startswith('### '):
                    # Sub-sub heading
                    doc.add_heading(line[4:], level=4)
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or any(line.startswith(f'{i}. ') for i in range(1, 10)):
                    # Numbered list
                    doc.add_paragraph(line[3:], style='List Number')
                elif line:
                    # Regular paragraph
                    doc.add_paragraph(line)
                else:
                    # Empty line - add space
                    doc.add_paragraph()
        
        doc.add_page_break()
    
    # Code Structure Section
    structure_heading = doc.add_heading('Code Structure', level=1)
    structure_heading.runs[0].font.color.rgb = csu_brown
    
    doc.add_paragraph('File and Directory Organization:')
    
    # Create table for file structure
    structure_table = doc.add_table(rows=1, cols=2)
    structure_table.style = 'Table Grid'
    
    # Header row
    header_cells = structure_table.rows[0].cells
    header_cells[0].text = 'File/Directory'
    header_cells[1].text = 'Purpose'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = csu_brown
    
    # File structure data
    structure_data = [
        ("app.py", "Flask application configuration and database setup"),
        ("main.py", "Application entry point"),
        ("models.py", "SQLAlchemy database models (User, Post, Comment)"),
        ("routes.py", "URL routing and view functions"),
        ("forms.py", "WTForms form definitions and validation"),
        ("utils.py", "Utility functions and decorators"),
        ("templates/", "Jinja2 HTML templates"),
        ("static/css/", "Custom CSS styling with CSU CCIS theme"),
        ("static/js/", "JavaScript for client-side functionality"),
        ("static/images/", "CSU CCIS logo and other images"),
        ("docs/", "Project documentation files")
    ]
    
    # Add data rows
    for file_name, purpose in structure_data:
        row_cells = structure_table.add_row().cells
        row_cells[0].text = file_name
        row_cells[1].text = purpose
    
    doc.add_paragraph()
    
    # Technologies Used Section
    tech_heading = doc.add_heading('Technologies Used', level=1)
    tech_heading.runs[0].font.color.rgb = csu_brown
    
    # Create table for technologies
    tech_table = doc.add_table(rows=1, cols=4)
    tech_table.style = 'Table Grid'
    
    # Header row
    tech_header_cells = tech_table.rows[0].cells
    tech_header_cells[0].text = 'Category'
    tech_header_cells[1].text = 'Technology'
    tech_header_cells[2].text = 'Version'
    tech_header_cells[3].text = 'Purpose'
    
    # Make header bold
    for cell in tech_header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = csu_brown
    
    # Technology data
    tech_data = [
        ("Backend", "Flask", "Latest", "Web framework"),
        ("Database", "PostgreSQL", "Latest", "Primary database"),
        ("ORM", "SQLAlchemy", "Latest", "Database abstraction"),
        ("Authentication", "Flask-Login", "Latest", "Session management"),
        ("Forms", "Flask-WTF", "Latest", "Form handling and validation"),
        ("Frontend", "Bootstrap 5", "5.3.0", "CSS framework"),
        ("JavaScript", "Vanilla JS", "ES6+", "Client-side functionality"),
        ("Server", "Gunicorn", "Latest", "WSGI HTTP server"),
        ("Security", "Werkzeug", "Latest", "Password hashing")
    ]
    
    # Add technology rows
    for category, technology, version, purpose in tech_data:
        tech_row_cells = tech_table.add_row().cells
        tech_row_cells[0].text = category
        tech_row_cells[1].text = technology
        tech_row_cells[2].text = version
        tech_row_cells[3].text = purpose
    
    # Add final sections
    doc.add_paragraph()
    
    features_heading = doc.add_heading('Key Features Implemented', level=1)
    features_heading.runs[0].font.color.rgb = csu_brown
    
    features = [
        "User Registration and Authentication with secure password hashing",
        "Blog Post Management (Create, Read, Update, Delete)",
        "Comment System for user interaction",
        "Responsive Design with CSU CCIS branding",
        "User Dashboard for content management",
        "Form Validation and Error Handling",
        "Session Management and Security",
        "Database Relationships and Data Integrity",
        "RESTful API Endpoints",
        "Mobile-Friendly Interface"
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Assignment requirements section
    doc.add_page_break()
    
    requirements_heading = doc.add_heading('Assignment Requirements Met', level=1)
    requirements_heading.runs[0].font.color.rgb = csu_brown
    
    doc.add_paragraph('This project fulfills all requirements for the 21-day course assignment:')
    
    requirements = [
        "✓ Simple web application using Flask framework",
        "✓ User authentication system (sign-up, sign-in, sign-out)",
        "✓ Two CRUD features: Blog Posts and Comments",
        "✓ Frontend-Backend-Database integration",
        "✓ CSU CCIS branding and professional design",
        "✓ Comprehensive documentation package",
        "✓ Database design with proper relationships",
        "✓ Security best practices implementation",
        "✓ Responsive and accessible user interface",
        "✓ Production-ready deployment configuration"
    ]
    
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    # Save document
    try:
        doc.save('output/CSU_CCIS_Blog_Documentation.docx')
        print("✅ Word documentation generated successfully!")
        print("📝 File saved as: output/CSU_CCIS_Blog_Documentation.docx")
        return True
    except Exception as e:
        print(f"❌ Error generating Word document: {e}")
        return False

if __name__ == "__main__":
    print("🔄 Generating CSU CCIS Blog Word Documentation...")
    create_word_documentation()